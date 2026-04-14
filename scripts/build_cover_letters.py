#!/usr/bin/env python3
"""
Build cover_letter.docx from cover_letter.json under /output/.

Usage (from repo root):
  pip install python-docx
  python scripts/build_cover_letters.py
  python scripts/build_cover_letters.py --skip-existing
"""

from __future__ import annotations

import argparse
import json
import os
import sys
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_DIR = REPO_ROOT / "output"

# Fixed letterhead (edit to match your résumé / preferences)
FULL_NAME = "Manas Dalvi"
ADDRESS_LINE = "Los Angeles, United States"
PHONE = "(213) 803-6398"
EMAIL = "manasman@usc.edu"
LINKEDIN = "linkedin.com/in/manasdalvi"
WEBSITE = "manasdalvi.vercel.app"


def _require_docx():
    try:
        from docx import Document  # noqa: F401
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt

        return Document, WD_ALIGN_PARAGRAPH, Inches, Pt
    except ImportError:
        print("❌ python-docx not installed. Run: pip install python-docx")
        sys.exit(1)


def _set_narrow_margins(section, Inches):
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)


def _add_center_heading(doc, WD_ALIGN_PARAGRAPH, Pt, text: str, *, size_pt: int, bold: bool) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size_pt)


def _add_center_line(doc, WD_ALIGN_PARAGRAPH, Pt, text: str, *, size_pt: int = 10) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size_pt)


def _add_contact_line(doc, WD_ALIGN_PARAGRAPH, Pt):
    """Plain underlined contact segments (readable in Word without field codes)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    sep = " | "
    size = Pt(10)

    p.add_run(PHONE).font.size = size
    for label in (EMAIL, LINKEDIN, WEBSITE):
        p.add_run(sep).font.size = size
        r = p.add_run(label)
        r.font.size = size
        r.font.underline = True


def build_docx(data: dict, out_path: str) -> None:
    Document, WD_ALIGN_PARAGRAPH, Inches, Pt = _require_docx()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    _set_narrow_margins(doc.sections[0], Inches)

    _add_center_heading(doc, WD_ALIGN_PARAGRAPH, Pt, FULL_NAME, size_pt=22, bold=True)
    _add_center_line(doc, WD_ALIGN_PARAGRAPH, Pt, ADDRESS_LINE, size_pt=10)
    _add_contact_line(doc, WD_ALIGN_PARAGRAPH, Pt)

    doc.add_paragraph()

    date_str = (data.get("date") or "").strip()
    if date_str:
        p = doc.add_paragraph(date_str)
        p.paragraph_format.space_after = Pt(12)

    hm = (data.get("hiring_manager") or "Hiring Manager").strip()
    co = (data.get("company") or "").strip()
    loc = (data.get("location") or "").strip()

    block = doc.add_paragraph()
    block.add_run(hm)
    block.add_run("\n")
    block.add_run(co)
    if loc:
        block.add_run("\n")
        block.add_run(loc)
    block.paragraph_format.space_after = Pt(12)

    sal = (data.get("salutation") or "Dear Hiring Manager,").strip()
    s = doc.add_paragraph(sal)
    s.paragraph_format.space_after = Pt(8)

    for key in ("opening", "experience", "closing"):
        text = (data.get(key) or "").strip()
        if not text:
            continue
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()
    close = doc.add_paragraph("Sincerely,")
    close.paragraph_format.space_after = Pt(4)

    sig = doc.add_paragraph()
    sig.add_run(FULL_NAME).bold = True

    doc.save(out_path)


def load_json(path: str) -> dict:
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--skip-existing", action="store_true")
    args = parser.parse_args()

    _require_docx()

    json_files = []
    for root, _, files in os.walk(str(OUTPUT_DIR)):
        for f in files:
            if f == "cover_letter.json":
                json_files.append(os.path.join(root, f))

    if not json_files:
        print(f"No cover_letter.json files found in {OUTPUT_DIR.relative_to(REPO_ROOT)}/")
        print("Run opencode with prompts/cover_letter_task.md first.")
        return

    print(f"Found {len(json_files)} cover_letter.json files\n")

    success = skipped = failed = 0

    for i, json_path in enumerate(sorted(json_files), 1):
        docx_path = os.path.join(os.path.dirname(json_path), "cover_letter.docx")
        label = os.path.dirname(json_path).replace(str(OUTPUT_DIR) + os.sep, "")

        if args.skip_existing and os.path.exists(docx_path):
            print(f"[{i:>3}/{len(json_files)}] ⏭  {label}")
            skipped += 1
            continue

        print(f"[{i:>3}/{len(json_files)}] 🔄  {label}")
        try:
            data = load_json(json_path)
            build_docx(data, docx_path)
            print(f"            ✅ cover_letter.docx saved")
            success += 1
        except (json.JSONDecodeError, OSError, ValueError) as e:
            print(f"            ❌ {e}")
            failed += 1

    print(f"\n{'─'*50}")
    print(f"Done.  ✅ {success}  ⏭ {skipped}  ❌ {failed}")


if __name__ == "__main__":
    main()
